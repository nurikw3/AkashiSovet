"""Генерация пояснительной записки в формате DOCX."""

from __future__ import annotations

import asyncio
import hashlib
import io
import json
from datetime import datetime
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from bot.logger import logger
import stdlib.db as db
import stdlib.redis_client as redis_client_module
import stdlib.s3 as s3
from stdlib.document_common import (
    build_sections,
    normalize_user_text,
    parse_attachments_field,
    signer_name_from_data,
    signer_position_from_data,
    split_section_body,
    topic_from_data,
)
from stdlib.template import (
    DOCX_TEMPLATE_REVISION_KEY,
    ApplicationTemplate,
    get_template,
)
from stdlib.text_normalize import attachment_name_for_document
from stdlib.timezone_util import ensure_app_tz, format_app_date_only

ASSETS_DIR = Path(__file__).parent / "assets"
TEMPLATE_PATH = ASSETS_DIR / "board_memo_template.docx"
# Футер в колонтитуле шаблона: word/media/image2.png (копия image2_updated.png).
FOOTER_IMAGE_PATH = ASSETS_DIR / "image2_updated.png"

DOCX_CACHE_KEY_FMT = "docx_cache:{app_id}"
DOCX_CACHE_TTL_SEC = 7 * 24 * 3600

_FONT_NAME = "Times New Roman"
_FONT_SIZE = Pt(12)

# Поля как в прежнем PDF-генераторе (ReportLab).
_MARGIN_TOP = Cm(2)
_MARGIN_BOTTOM = Cm(2)
_MARGIN_LEFT = Cm(3)
_MARGIN_RIGHT = Cm(1.5)
_BODY_FIRST_LINE = Cm(1.25)
_LIST_LEFT = Cm(1.25)
_LIST_HANGING = Cm(0.63)


def _docx_cache_token(
    app_raw: dict,
    full_name: str | None,
    position: str | None,
    sig_ref: str | None,
    tpl: ApplicationTemplate,
    template_revision: str,
) -> str:
    blocks = app_raw.get("blocks") or ""
    attachments = app_raw.get("attachments") or ""
    tpl_json = json.dumps(
        [b.model_dump(mode="json") for b in tpl.blocks],
        sort_keys=True,
        ensure_ascii=False,
    )
    footer_rev = ""
    if FOOTER_IMAGE_PATH.exists():
        footer_rev = hashlib.sha256(FOOTER_IMAGE_PATH.read_bytes()).hexdigest()[:16]
    parts = [
        str(blocks),
        str(attachments),
        full_name or "",
        position or "",
        sig_ref or "",
        tpl_json,
        template_revision,
        footer_rev,
        "docx-v3",
    ]
    return hashlib.sha256("|".join(parts).encode("utf-8")).hexdigest()


async def _get_docx_template_revision() -> str:
    r = redis_client_module.redis_client
    if not r:
        return "0"
    try:
        v = await r.get(DOCX_TEMPLATE_REVISION_KEY)
        return v if v is not None else "0"
    except Exception:
        return "0"


_TEMPLATE_FOOTER_MEDIA = "word/media/image2.png"
_template_footer_digest: str | None = None


def _ensure_template_footer_image() -> None:
    """Встраивает image2_updated.png в колонтитул шаблона, если файл изменился."""
    global _template_footer_digest
    if not FOOTER_IMAGE_PATH.exists() or not TEMPLATE_PATH.exists():
        return
    source_digest = hashlib.sha256(FOOTER_IMAGE_PATH.read_bytes()).hexdigest()
    if _template_footer_digest == source_digest:
        return
    with ZipFile(TEMPLATE_PATH, "r") as zin:
        try:
            embedded = zin.read(_TEMPLATE_FOOTER_MEDIA)
        except KeyError:
            embedded = b""
        if hashlib.sha256(embedded).hexdigest() == source_digest:
            _template_footer_digest = source_digest
            return
        buf = io.BytesIO()
        with ZipFile(buf, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == _TEMPLATE_FOOTER_MEDIA:
                    data = FOOTER_IMAGE_PATH.read_bytes()
                zout.writestr(item, data)
    TEMPLATE_PATH.write_bytes(buf.getvalue())
    _template_footer_digest = source_digest
    logger.debug("DOCX template footer synced from {}", FOOTER_IMAGE_PATH.name)


def _clear_body_keep_sectpr(doc: Document) -> None:
    body = doc.element.body
    for element in list(body):
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def _apply_page_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = _MARGIN_TOP
        section.bottom_margin = _MARGIN_BOTTOM
        section.left_margin = _MARGIN_LEFT
        section.right_margin = _MARGIN_RIGHT


def _apply_run_font(run, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = _FONT_NAME
    run.font.size = _FONT_SIZE
    run.bold = bold
    run.italic = italic


def _add_styled_paragraph(
    doc: Document,
    text: str,
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool = False,
    italic: bool = False,
    space_before: Pt | None = None,
    space_after: Pt | None = None,
) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.paragraph_format.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    _apply_run_font(run, bold=bold, italic=italic)


def _add_body_paragraph(doc: Document, text: str, *, bullet: bool = False) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)
    if bullet:
        pf.left_indent = _LIST_LEFT
        pf.first_line_indent = -_LIST_HANGING
    else:
        pf.first_line_indent = _BODY_FIRST_LINE
    prefix = "• " if bullet else ""
    run = p.add_run(f"{prefix}{text}")
    _apply_run_font(run)


def _add_section(doc: Document, title: str, body: str) -> None:
    title = normalize_user_text(title)
    _add_styled_paragraph(
        doc,
        title,
        bold=True,
        space_before=Pt(8),
        space_after=Pt(4),
    )
    lines, has_numbering = split_section_body(body)
    for line in lines:
        _add_body_paragraph(doc, line, bullet=has_numbering)


def _generate_docx_sync(
    data: dict,
    tpl: ApplicationTemplate | None,
    blocks_map: dict[str, str] | None,
    signature_bytes: bytes | None,
) -> BytesIO:
    if not data:
        data = {}

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"DOCX template not found: {TEMPLATE_PATH}")

    _ensure_template_footer_image()
    doc = Document(str(TEMPLATE_PATH))
    _clear_body_keep_sectpr(doc)
    _apply_page_margins(doc)

    _add_styled_paragraph(
        doc, "Членам Правления", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True
    )
    _add_styled_paragraph(
        doc,
        "ПК «AKASHI Data Center PLC»",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        bold=True,
        space_after=Pt(10),
    )

    _add_styled_paragraph(
        doc,
        "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        space_before=Pt(10),
        space_after=Pt(4),
    )
    topic = normalize_user_text(topic_from_data(data, tpl, blocks_map))
    _add_styled_paragraph(
        doc,
        f"по вопросу «{topic}»",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(10),
    )

    sections, attach_num = build_sections(data, tpl, blocks_map)
    for sec_title, body in sections:
        _add_section(doc, sec_title, body)

    attach_title = f"{attach_num}. Приложения / дополнительные материалы:"
    _add_styled_paragraph(
        doc,
        attach_title,
        bold=True,
        space_before=Pt(8),
        space_after=Pt(4),
    )

    attachments = parse_attachments_field(data.get("attachments"))
    names = [
        normalize_user_text(attachment_name_for_document(str(name)))
        for name in attachments
        if isinstance(attachments, list) and str(name).strip()
    ]
    if names:
        for name in names:
            _add_body_paragraph(doc, name, bullet=True)
    else:
        _add_body_paragraph(doc, "(приложения не прикреплены)")

    doc.add_paragraph()
    position = signer_position_from_data(data)
    username = signer_name_from_data(data)
    date = str(data.get("date") or "__.__.____")

    p_pos = doc.add_paragraph()
    p_pos.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_pos = p_pos.add_run(f"{position}:")
    _apply_run_font(run_pos, bold=True)

    if signature_bytes:
        try:
            sig_stream = io.BytesIO(signature_bytes)
            pic_p = doc.add_paragraph()
            pic_p.paragraph_format.space_after = Pt(4)
            pic_p.add_run().add_picture(sig_stream, width=Cm(4), height=Cm(2))
        except Exception as e:
            logger.warning("DOCX: ошибка вставки подписи: {}", e)
            doc.add_paragraph()
    else:
        doc.add_paragraph()

    _add_body_paragraph(doc, f"________________ /{username}/")
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_before = Pt(10)
    run_label = p_date.add_run("Дата: ")
    _apply_run_font(run_label, bold=True)
    run_val = p_date.add_run(date)
    _apply_run_font(run_val)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


async def _load_signature_for_ref(ref: str | bytes | bytearray | None) -> bytes | None:
    if not ref:
        return None
    if isinstance(ref, (bytes, bytearray)):
        return bytes(ref)
    if not isinstance(ref, str):
        return None
    if ref.startswith("signatures/"):
        try:
            data = await s3.download_bytes(ref, s3.BUCKET_SIGNATURES)
            return data if data else None
        except RuntimeError:
            logger.warning("DOCX: S3 не настроен — подпись недоступна")
            return None
        except Exception as e:
            logger.warning("DOCX: не удалось скачать подпись из S3: {}", e)
            return None
    return None


async def _resolve_user_signature_bytes(user_id: int) -> bytes | None:
    ref = await db.get_user_signature(user_id)
    return await _load_signature_for_ref(ref)


async def generate_docx(
    data: dict,
    user_id: int | None = None,
    signature_bytes: bytes | None = None,
    *,
    tpl: ApplicationTemplate | None = None,
    blocks_map: dict[str, str] | None = None,
) -> BytesIO:
    try:
        if not data:
            data = {}

        if signature_bytes is None and user_id:
            try:
                signature_bytes = await _resolve_user_signature_bytes(user_id)
            except Exception:
                signature_bytes = None

        buffer = await asyncio.to_thread(
            _generate_docx_sync, data, tpl, blocks_map, signature_bytes
        )
        logger.info("DOCX generated successfully for app_id={}", data.get("app_id"))
        return buffer
    except Exception as e:
        logger.error("DOCX generation failed for app_id={}: {}", data.get("app_id"), e)
        import traceback

        logger.error("Traceback: {}", traceback.format_exc())
        raise


def generate_docx_filename(
    full_name: str | None, position: str | None, dt: datetime
) -> str:
    safe_name = (full_name or "Сотрудник").replace(" ", "_")
    safe_pos = (position or "Должность").replace(" ", "_")
    time_str = ensure_app_tz(dt).strftime("%H-%M_%d-%m-%Y")
    return f"{time_str}_{safe_name}_{safe_pos}.docx"


def resolve_application_docx_filename(
    app_row: dict,
    *,
    full_name: str | None,
    position: str | None,
    dt: datetime | None,
) -> str:
    preferred = str(app_row.get("main_pdf_filename") or "").strip()
    if preferred:
        safe = Path(preferred).name
        lower = safe.lower()
        if lower.endswith(".docx") or lower.endswith(".pdf"):
            return safe
        return f"{safe}.docx"
    base_dt = dt or datetime.now()
    return generate_docx_filename(full_name, position, base_dt)


def media_type_for_filename(filename: str) -> str:
    if filename.lower().endswith(".pdf"):
        return "application/pdf"
    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


async def invalidate_docx_cache(app_id: int, *, user_id: int | None = None) -> None:
    await invalidate_docx_content_cache(app_id, user_id=user_id)


async def invalidate_docx_delivery_cache(app_id: int) -> None:
    try:
        await db.set_pdf_file_id(app_id, None)
    except Exception as e:
        logger.warning(
            "invalidate_docx_delivery_cache clear pdf_file_id app_id={}: {}",
            app_id,
            e,
        )


async def invalidate_docx_content_cache(app_id: int, *, user_id: int | None = None) -> None:
    r = redis_client_module.redis_client
    cache_key = DOCX_CACHE_KEY_FMT.format(app_id=app_id)
    if r:
        try:
            await r.delete(cache_key)
        except Exception as e:
            logger.warning("invalidate_docx_cache redis delete app_id={}: {}", app_id, e)

    await invalidate_docx_delivery_cache(app_id)

    if not s3.is_s3_configured():
        return
    uid = user_id
    if uid is None:
        row = await db.get_app(app_id)
        if not row:
            return
        uid = row["user_id"]
    for key_fn in (s3.docx_key, s3.pdf_key):
        key = key_fn(uid, app_id)
        try:
            await s3.delete_object(key, s3.BUCKET_PDF)
        except Exception as e:
            logger.warning(
                "invalidate_docx_cache s3 delete app_id={} key={}: {}",
                app_id,
                key,
                e,
            )


async def get_app_docx_buffer(app_id: int) -> BytesIO:
    app_raw = await db.get_app(app_id)
    if not app_raw:
        raise ValueError(f"App {app_id} not found")

    uploaded_key = app_raw.get("main_pdf_s3_key")
    if uploaded_key:
        try:
            uploaded_bytes = await s3.download_bytes(uploaded_key, s3.BUCKET_PDF)
            if uploaded_bytes:
                logger.debug(
                    "Primary uploaded document used | app_id={} size_bytes={}",
                    app_id,
                    len(uploaded_bytes),
                )
                return BytesIO(uploaded_bytes)
            logger.warning(
                "Uploaded document key exists but object missing | app_id={} key={}",
                app_id,
                uploaded_key,
            )
        except Exception as e:
            logger.warning(
                "Failed to load uploaded main document; falling back to generated | app_id={} err={}",
                app_id,
                e,
            )

    u_id = app_raw["user_id"]
    try:
        blocks = json.loads(app_raw.get("blocks", "{}"))
    except Exception:
        blocks = {}

    tpl, full_name, position, sig_ref, tpl_rev = await asyncio.gather(
        get_template(),
        db.get_user_full_name(u_id),
        db.get_user_position(u_id),
        db.get_user_signature(u_id),
        _get_docx_template_revision(),
    )
    token = _docx_cache_token(app_raw, full_name, position, sig_ref, tpl, tpl_rev)

    r = redis_client_module.redis_client
    cache_key = DOCX_CACHE_KEY_FMT.format(app_id=app_id)
    if r and s3.is_s3_configured():
        try:
            cached = await r.get(cache_key)
            if cached == token:
                docx_key = s3.docx_key(u_id, app_id)
                docx_bytes = await s3.download_bytes(docx_key, s3.BUCKET_PDF)
                if docx_bytes:
                    logger.debug(
                        "DOCX cache hit app_id={} size_bytes={}",
                        app_id,
                        len(docx_bytes),
                    )
                    return BytesIO(docx_bytes)
                logger.debug("DOCX cache miss app_id={} reason=s3_missing", app_id)
            elif cached:
                logger.debug("DOCX cache miss app_id={} reason=token_mismatch", app_id)
            else:
                logger.debug("DOCX cache miss app_id={} reason=redis_empty", app_id)
        except Exception as e:
            logger.warning("DOCX cache read failed app_id={}: {}", app_id, e)

    await invalidate_docx_delivery_cache(app_id)

    raw_att = app_raw.get("attachments")
    parsed_atts = parse_attachments_field(raw_att)
    clean_atts = []
    if isinstance(parsed_atts, list):
        for f in parsed_atts:
            if isinstance(f, dict):
                raw_name = f.get("name") or f.get("file_name") or "Файл"
                clean_atts.append(attachment_name_for_document(str(raw_name)))
            else:
                clean_atts.append(attachment_name_for_document(str(f)))

    docx_data = {
        "app_id": app_id,
        "attachments": clean_atts,
        "full_name": full_name,
        "position": position,
        "date": format_app_date_only(app_raw["created_at"]),
    }

    signature_bytes = await _load_signature_for_ref(sig_ref)
    buf = await generate_docx(
        docx_data,
        user_id=None,
        signature_bytes=signature_bytes,
        tpl=tpl,
        blocks_map=blocks,
    )

    if r and s3.is_s3_configured():
        try:
            body = buf.getvalue()
            logger.info(
                "DOCX ready | app_id={} size_bytes={} size_kb={:.1f}",
                app_id,
                len(body),
                len(body) / 1024,
            )
            await s3.upload_bytes(
                body,
                s3.docx_key(u_id, app_id),
                s3.BUCKET_PDF,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            await r.set(cache_key, token, ex=DOCX_CACHE_TTL_SEC)
        except Exception as e:
            logger.warning("DOCX cache store failed app_id={}: {}", app_id, e)
    else:
        logger.info(
            "DOCX ready | app_id={} size_bytes={} size_kb={:.1f}",
            app_id,
            len(buf.getbuffer()),
            len(buf.getbuffer()) / 1024,
        )

    buf.seek(0)
    return buf
